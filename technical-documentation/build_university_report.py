from __future__ import annotations

import os
import sys
from io import BytesIO
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
FIGURE_DIR = ROOT / "screenshots-annotated"
OUTPUT = ROOT / "VolUnite_Technical_Documentation.docx"
SKILL_DIR = Path(
    "/Users/manuelv/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.709.11516/skills/documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 40, 58)
MUTED = RGBColor(92, 98, 112)
LIGHT_FILL = "F2F4F7"
RED = "D71920"
WHITE = RGBColor(255, 255, 255)
CALLOUT_TEXT = RGBColor(50, 55, 66)


FIGURES = [
    {
        "heading": "2. Authentication and Public Discovery",
        "image": "01-login.jpg",
        "caption": "Figure 1. Login and public organisation discovery screen.",
        "summary": (
            "The landing screen combines authentication with limited public discovery. "
            "A visitor can review organisations before signing in, while authenticated "
            "access is initiated through a conventional username and password workflow."
        ),
        "items": [
            "The page header communicates the current route and exposes the session action.",
            "The authentication form provides username, password, password visibility, remember-me and recovery controls.",
            "Organisation search and filtering are available without authentication.",
            "Organisation cards present the name, short purpose and thematic tags.",
        ],
    },
    {
        "heading": "3. Personalised Home Dashboard",
        "image": "03-home.jpg",
        "caption": "Figure 2. Authenticated home dashboard for the demo user Alice.",
        "summary": (
            "After authentication, the application presents a personalised operational "
            "summary. The dashboard prioritises recent participation and recommends "
            "new opportunities using the user's existing profile and activity history."
        ),
        "items": [
            "The user summary displays identity, membership date and totals for friends, organisations and activities.",
            "Recent highlights show the latest joined activity, organisation and friend.",
            "Recommended activity cards expose date, location, tags, skills, capacity and a direct join action.",
            "A persistent bottom navigation bar provides access to the main application modules.",
        ],
    },
    {
        "heading": "4. Organisation Management",
        "image": "05-organisation-detail.jpg",
        "caption": "Figure 3. Organisation detail page with operational and social information.",
        "summary": (
            "Organisation pages consolidate descriptive, geographic and operational data. "
            "They connect community goals with concrete projects, activities and membership "
            "information, allowing volunteers to evaluate and join an organisation."
        ),
        "items": [
            "The organisation identity block includes its description, creation date, location and join action.",
            "A map visualises the organisation's geographic context.",
            "Community goals show measurable progress and provide create, overview and edit actions.",
            "Projects are presented with status information.",
            "Activities expose schedule, capacity and participant information.",
            "Membership sections distinguish friends in the organisation from the complete member list.",
        ],
    },
    {
        "heading": "5. Community and Collaboration",
        "image": "06a-community-friends.jpg",
        "caption": "Figure 4. Community screen showing friendships, shared activities and relationship analysis.",
        "summary": (
            "The community module extends the platform beyond activity registration. "
            "It combines social relationships, shared participation, forum discussions, "
            "direct conversations and graph-based exploration in one tabbed workspace."
        ),
        "items": [
            "Top-level tabs separate friends, forum and chat functionality.",
            "The friends list links to individual volunteer profiles.",
            "Shared activity cards show where friends are participating.",
            "Secondary tabs switch between relationship and activity graphs.",
            "The relationship graph visualises the selected user's social network.",
            "The main navigation remains available during community interaction.",
        ],
    },
    {
        "heading": "6. Activity Discovery and Participation",
        "image": "11-activities.jpg",
        "caption": "Figure 5. Activity catalogue with joined, recommended and complete result groups.",
        "summary": (
            "Activity discovery is structured around user state. Joined activities are "
            "separated from recommendations and the complete catalogue, while search, "
            "filters, tags and capacity information support rapid comparison."
        ),
        "items": [
            "The overview displays total, joined and open activity counts.",
            "Search and filter controls narrow the activity catalogue.",
            "Joined activities provide a concise view of existing commitments.",
            "Recommended activities support personalised discovery.",
            "The complete catalogue presents all available opportunities with dates, tags and friend participation.",
            "Navigation and the create action support movement between browsing and administration workflows.",
        ],
    },
    {
        "heading": "7. Volunteer Profile",
        "image": "09-profile.jpg",
        "caption": "Figure 6. Volunteer profile and participation history.",
        "summary": (
            "The profile functions as both an identity record and a participation portfolio. "
            "It combines contact information, skills and interests with the user's social "
            "connections and volunteering history."
        ),
        "items": [
            "Profile identity and the edit action are displayed prominently.",
            "Account details, skills and fields of interest support matching and recommendations.",
            "Friends provide the social context of the volunteer.",
            "Joined organisations show institutional participation.",
            "Joined activities form a concise volunteering history.",
            "The profile can be shared and remains connected to the primary navigation.",
        ],
    },
    {
        "heading": "8. Reporting and Decision Support",
        "image": "17-reports.jpg",
        "caption": "Figure 7. Reports dashboard with operational and participation analytics.",
        "summary": (
            "The reporting module translates participation data into operational indicators. "
            "Filters and user selection support targeted analysis, while several chart types "
            "cover time, category, contribution and capacity perspectives."
        ),
        "items": [
            "Filters and the volunteer selector define the reporting scope.",
            "Key indicators summarise total hours, activities, active volunteers and average occupancy.",
            "The monthly trend chart shows activity volume over time.",
            "Category distribution highlights the balance of focus areas.",
            "Volunteer contribution compares participation between users.",
            "Capacity demand and the summary row identify highly requested activities and occupancy levels.",
        ],
    },
    {
        "heading": "9. Community Goals",
        "image": "18-community-goals.jpg",
        "caption": "Figure 8. Community-goal progress and contribution overview.",
        "summary": (
            "Community goals connect individual activity participation to measurable "
            "organisation-level outcomes. Each goal records a target, current progress, "
            "thematic tags and the completed activities that contributed to it."
        ),
        "items": [
            "The goals header provides direct access to goal creation.",
            "Each goal exposes title, tags, progress, target values and editing.",
            "Completed activity contributions can be expanded for traceability.",
            "Multiple goals can be monitored concurrently.",
            "The global navigation remains accessible from the goal workflow.",
        ],
    },
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color="D9DDE5", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("VOLUNITE  |  TECHNICAL UI DOCUMENTATION")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    label = footer_p.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(footer_p)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("VOLUNITE")
    set_run_font(run, size=13, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Technical User Interface Documentation")
    set_run_font(run, size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(52)
    run = p.add_run("Functional overview based on annotated application screenshots")
    set_run_font(run, size=14, color=MUTED)

    for label, value in (
        ("Document type", "University technical report"),
        ("System", "Volunteer management and community platform"),
        ("Prepared", "July 2026"),
        ("Student", "________________________________"),
        ("Course", "________________________________"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(56)
    run = p.add_run(
        "This document describes the principal user-facing functions visible in "
        "the captured frontend implementation."
    )
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def add_callout_table(doc, items):
    table = doc.add_table(rows=0, cols=2)
    for index, text in enumerate(items, start=1):
        cells = table.add_row().cells
        number_cell, text_cell = cells
        number_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(number_cell, RED)
        set_cell_fill(text_cell, LIGHT_FILL if index % 2 else "FFFFFF")
        set_cell_border(number_cell, color="FFFFFF", size="6")
        set_cell_border(text_cell, color="D9DDE5", size="4")

        p = number_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(index))
        set_run_font(run, size=10, color=WHITE, bold=True)

        p = text_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text)
        set_run_font(run, size=9.2, color=CALLOUT_TEXT)

    apply_table_geometry(
        table,
        [720, 8640],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 70, "bottom": 70, "start": 120, "end": 120},
    )
    return table


def add_scaled_image(
    doc, image_path, max_width=6.25, max_height=5.65, alt_text="Annotated application screenshot"
):
    with Image.open(image_path) as img:
        width_px, height_px = img.size
        normalized = img.convert("RGB")
        image_stream = BytesIO()
        normalized.save(image_stream, format="PNG", optimize=True)
        image_stream.seek(0)
    aspect = width_px / height_px
    width = min(max_width, max_height * aspect)
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(
        image_stream, width=Inches(width), height=Inches(height)
    )
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", alt_text)


def add_figure_section(doc, figure):
    doc.add_heading(figure["heading"], level=1)
    p = doc.add_paragraph(figure["summary"])
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    item_count = len(figure["items"])
    max_height = 5.10 if item_count >= 6 else 5.42 if item_count == 5 else 5.65
    add_scaled_image(
        doc,
        FIGURE_DIR / figure["image"],
        max_height=max_height,
        alt_text=figure["caption"],
    )
    caption = doc.add_paragraph(figure["caption"], style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout_table(doc, figure["items"])


def add_intro(doc):
    doc.add_page_break()
    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "VolUnite is a web-based volunteer coordination platform that connects "
        "volunteers with organisations, activities and other community members. "
        "This report gives a concise functional description of the user interface "
        "shown in the annotated screenshots. It focuses on observable system "
        "behaviour rather than implementation details that cannot be confirmed visually."
    )

    doc.add_heading("System purpose", level=2)
    doc.add_paragraph(
        "The platform supports the complete volunteer participation cycle: discovery, "
        "authentication, profile-based matching, activity registration, organisation "
        "membership, community communication, goal tracking and reporting."
    )

    doc.add_heading("Principal capabilities", level=2)
    capabilities = [
        "Public organisation discovery and authenticated access.",
        "Personalised dashboards and activity recommendations.",
        "Organisation, project, activity and community-goal management.",
        "Friendships, forums, direct chat and relationship graphs.",
        "Volunteer profiles containing skills, interests and participation history.",
        "Operational reporting through indicators, charts and capacity analysis.",
    ]
    add_callout_table(doc, capabilities)

    doc.add_heading("Interface structure", level=2)
    doc.add_paragraph(
        "Authenticated screens share a persistent header containing notification and "
        "session controls. A fixed bottom navigation bar links the primary modules: "
        "profile, community, home, organisations, activities and reports. Card-based "
        "layouts are used consistently to group related data and actions."
    )

    doc.add_heading("Evidence basis", level=2)
    doc.add_paragraph(
        "The figures use representative demo data for the user Alice. Red numbered "
        "boxes identify the interface regions discussed directly below each image. "
        "The screenshots document the current frontend state as captured in July 2026."
    )


def add_observations(doc):
    heading = doc.add_heading("10. Technical Observations and Conclusion", level=1)
    heading.paragraph_format.page_break_before = True

    doc.add_heading("Main strengths", level=2)
    doc.add_paragraph(
        "The interface uses consistent cards, navigation and search/filter patterns. "
        "Skills, interests, relationships and participation history support personalised "
        "recommendations, while community goals and reports convert individual activity "
        "records into measurable organisational outcomes."
    )

    doc.add_heading("Current interface defects", level=2)
    doc.add_paragraph(
        "The captured registration, project-detail and notifications routes show a "
        "collapsed route-container layout in which content is compressed into a narrow "
        "vertical column. This should be treated as a frontend layout defect rather than "
        "intended responsive behaviour."
    )

    doc.add_heading("Recommended technical priorities", level=2)
    priorities = [
        "Correct the route-container or component-host width rules on the affected pages.",
        "Add automated visual regression tests for desktop and mobile breakpoints.",
        "Verify that maps, graphs and charts provide accessible text alternatives.",
        "Apply consistent loading, empty and error states across all data-driven modules.",
    ]
    add_callout_table(doc, priorities)

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "VolUnite presents a coherent volunteer-management concept that combines opportunity "
        "discovery, social participation and measurable impact. The strongest parts of the "
        "interface are its personalised activity workflow, integrated community features and "
        "analytical reporting. Resolving the identified layout defects and strengthening "
        "accessibility would improve production readiness."
    )


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc)
    add_intro(doc)
    for figure in FIGURES:
        add_figure_section(doc, figure)
    add_observations(doc)

    props = doc.core_properties
    props.title = "VolUnite Technical User Interface Documentation"
    props.subject = "University technical report based on annotated screenshots"
    props.author = "Student submission"
    props.keywords = "VolUnite, volunteer management, technical documentation, UI"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
